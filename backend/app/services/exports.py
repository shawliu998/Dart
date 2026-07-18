"""Local, auditable export artifacts for the response workflow."""

from __future__ import annotations

from datetime import datetime
import hashlib
from pathlib import Path
from uuid import UUID

from docx import Document as WordDocument
from openpyxl import Workbook
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.auth.dependencies import Principal
from app.core.config import get_settings
from app.models.entities import (
    ComplianceCheck,
    EvidenceClaim,
    ResponseEvidenceLink,
    ResponseItem,
    Requirement,
    TenderProject,
)


def _export_root(project_id: UUID) -> tuple[Path, str]:
    settings = get_settings()
    root = settings.app_data_dir or Path(__file__).resolve().parents[2] / "data"
    relative = f"exports/{project_id}"
    directory = root / relative
    directory.mkdir(parents=True, exist_ok=True)
    return directory, relative


def _response_claims(db: Session, response_id: UUID) -> list[EvidenceClaim]:
    return list(
        db.scalars(
            select(EvidenceClaim)
            .join(ResponseEvidenceLink, ResponseEvidenceLink.evidence_claim_id == EvidenceClaim.id)
            .where(ResponseEvidenceLink.response_item_id == response_id)
        )
    )


def _responses(db: Session, principal: Principal, project_id: UUID) -> dict[UUID, ResponseItem]:
    return {
        item.requirement_id: item
        for item in db.scalars(
            select(ResponseItem).where(
                ResponseItem.project_id == project_id,
                ResponseItem.tenant_id == principal.tenant_id,
            )
        )
    }


def export_project_artifacts(db: Session, principal: Principal, project_id: UUID) -> list[dict]:
    """Create spreadsheet and Word deliverables from persisted, human-visible data."""
    project = db.scalar(
        select(TenderProject).where(
            TenderProject.id == project_id, TenderProject.tenant_id == principal.tenant_id
        )
    )
    if project is None:
        raise LookupError("project not found")
    requirements = list(
        db.scalars(
            select(Requirement)
            .where(
                Requirement.project_id == project_id,
                Requirement.tenant_id == principal.tenant_id,
                Requirement.is_current.is_(True),
            )
            .order_by(Requirement.source_page, Requirement.requirement_code)
        )
    )
    checks = {
        item.requirement_id: item
        for item in db.scalars(
            select(ComplianceCheck).where(
                ComplianceCheck.project_id == project_id,
                ComplianceCheck.tenant_id == principal.tenant_id,
            )
        )
        if item.requirement_id is not None
    }
    responses = _responses(db, principal, project_id)
    directory, relative = _export_root(project_id)

    matrix_name = "合规矩阵.xlsx"
    matrix_path = directory / matrix_name
    workbook = Workbook()
    sheet = workbook.active
    if sheet is None:
        raise RuntimeError("XLSX workbook did not create an active sheet")
    sheet.title = "合规矩阵"
    sheet.append(["编号", "分类", "风险", "招标原文", "结构化要求", "来源页码", "是否强制", "否决候选", "关联证据（可能尚未人工确认）", "合规结果", "响应状态", "响应内容", "缺失信息"])
    for requirement in requirements:
        response = responses.get(requirement.id)
        claims = _response_claims(db, response.id) if response else []
        check = checks.get(requirement.id)
        sheet.append([
            requirement.requirement_code or "",
            requirement.category,
            requirement.risk_level,
            requirement.original_text,
            requirement.normalized_requirement,
            requirement.source_page,
            "是" if requirement.mandatory else "否",
            "是" if requirement.disqualification_if_failed else "否",
            "；".join(f"{claim.subject}：{claim.value}" for claim in claims),
            check.result if check else "manual_review",
            response.status if response else "not_started",
            (response.edited_text or response.draft_text) if response else "【待编写】",
            "；".join(response.missing_information) if response else "",
        ])
    workbook.save(matrix_path)

    draft_name = "投标响应初稿.docx"
    draft_path = directory / draft_name
    document = WordDocument()
    document.add_heading(project.name, level=0)
    document.add_paragraph(f"项目编号：{project.project_code}")
    document.add_paragraph(f"采购人：{project.buyer_name}")
    document.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    document.add_paragraph("内部草稿：要求、证据与响应可能包含尚未人工确认的暂定结果，不得直接用于外部提交。")
    for requirement in requirements:
        response = responses.get(requirement.id)
        document.add_heading(f"{requirement.requirement_code or ''} {requirement.title}", level=1)
        document.add_paragraph(f"招标要求：{requirement.normalized_requirement}")
        document.add_paragraph(f"响应内容：{(response.edited_text or response.draft_text) if response else '【待编写】'}")
        claims = _response_claims(db, response.id) if response else []
        document.add_paragraph(f"引用材料：{'；'.join(f'{claim.subject}：{claim.value}' for claim in claims) or '【待补充】'}")
        if response and response.missing_information:
            document.add_paragraph(f"待补内容：{'；'.join(response.missing_information)}")
    document.save(str(draft_path))

    risks_name = "风险与待办.xlsx"
    risks_path = directory / risks_name
    risk_book = Workbook()
    risk_sheet = risk_book.active
    if risk_sheet is None:
        raise RuntimeError("risk workbook did not create an active sheet")
    risk_sheet.title = "风险与待办"
    risk_sheet.append(["风险类型", "风险等级", "问题描述", "来源", "建议动作", "状态"])
    for requirement in requirements:
        response = responses.get(requirement.id)
        check = checks.get(requirement.id)
        if (check and check.result in {"fail", "warning", "manual_review"}) or (response and response.status == "missing_evidence"):
            risk_sheet.append([
                "合规检查" if check else "响应草稿",
                requirement.risk_level,
                (check.reason if check else "缺少人工接受的证据"),
                f"第 {requirement.source_page} 页",
                "补充材料或人工复核",
                check.result if check else (response.status if response else "missing_evidence"),
            ])
    risk_book.save(risks_path)

    return [
        {"artifact_type": "compliance_matrix_xlsx", "title": matrix_name, "storage_key": f"{relative}/{matrix_name}", "content_hash": hashlib.sha256(matrix_path.read_bytes()).hexdigest(), "metadata": {"download_name": matrix_name}},
        {"artifact_type": "response_draft_docx", "title": draft_name, "storage_key": f"{relative}/{draft_name}", "content_hash": hashlib.sha256(draft_path.read_bytes()).hexdigest(), "metadata": {"download_name": draft_name}},
        {"artifact_type": "risk_tasks_xlsx", "title": risks_name, "storage_key": f"{relative}/{risks_name}", "content_hash": hashlib.sha256(risks_path.read_bytes()).hexdigest(), "metadata": {"download_name": risks_name}},
    ]
