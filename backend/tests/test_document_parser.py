from __future__ import annotations

from io import BytesIO

from docx import Document
from openpyxl import Workbook
from pypdf import PdfWriter

from app.parsers.deterministic import DeterministicTextParser


def test_xlsx_parser_keeps_sheet_name_cell_addresses_and_formulas() -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "报价表"
    sheet["A1"] = "项目"
    sheet["B2"] = "=1+1"
    stream = BytesIO()
    workbook.save(stream)

    pages = DeterministicTextParser().parse(
        stream.getvalue(),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

    assert pages[0].layout_json["sheet"] == "报价表"
    assert {cell["address"] for cell in pages[0].layout_json["cells"]} == {"A1", "B2"}
    formula = next(cell for cell in pages[0].layout_json["cells"] if cell["address"] == "B2")
    assert formula["cached_value"] is None
    assert formula["needs_review"] is True
    assert "B2: =1+1 [公式缓存值缺失，需人工复核]" in pages[0].raw_text


def test_blank_pdf_page_is_marked_for_ocr_without_claiming_ocr_was_run() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    stream = BytesIO()
    writer.write(stream)

    pages = DeterministicTextParser().parse(stream.getvalue(), "application/pdf")

    assert pages[0].raw_text == ""
    assert pages[0].layout_json["ocr_required"] is True
    assert pages[0].ocr_used is False


def test_docx_parser_keeps_logical_paragraph_and_table_locations() -> None:
    document = Document()
    document.add_paragraph("资格要求")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "证书"
    table.cell(0, 1).text = "ISO9001"
    document.add_paragraph("交付要求")
    stream = BytesIO()
    document.save(stream)

    pages = DeterministicTextParser().parse(
        stream.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert pages[0].layout_json["blocks"][0]["paragraph_index"] == 1
    assert [block["type"] for block in pages[0].layout_json["blocks"]] == [
        "paragraph",
        "table",
        "paragraph",
    ]
    assert pages[0].layout_json["tables"][0]["rows"][0][1]["text"] == "ISO9001"
    assert pages[0].raw_text.splitlines() == ["资格要求", "证书 | ISO9001", "交付要求"]


def test_image_is_routed_to_ocr_without_decoding_binary_as_text() -> None:
    pages = DeterministicTextParser().parse(b"\x89PNG\r\n\x1a\n", "image/png")

    assert pages[0].raw_text == ""
    assert pages[0].layout_json["ocr_required"] is True
    assert pages[0].ocr_used is False
